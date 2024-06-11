import os

from docx import Document
import matplotlib.pyplot as plt
import io

from datetime import datetime


class ReportGenerator:

    def __init__(self):
        self.progression = {'excellent': ['#31C013', '#B5EFA9'],
                            'good': ['#68CC14', '#CCF2AB'],
                            'normal': ['#E6E317', '#F9F8B0'],
                            'bad': ['#E61C17', '#F9B2B0',]}
        current_dir = os.path.dirname(os.path.abspath(__file__))
        self.template_path = os.path.join(current_dir, "template/template.docx")

    def _load_template(self):
        doc = Document(self.template_path)
        return doc

    def _generate_progress_bar(self, plagiarism_percentage):
        if plagiarism_percentage >= 90:
            colors = self.progression.get('excellent')
            short_result = 'Excellent'
        elif plagiarism_percentage >= 80:
            colors = self.progression.get('good')
            short_result = 'Good'
        elif plagiarism_percentage >= 70:
            colors = self.progression.get('normal')
            short_result = 'Normal'
        else:
            colors = self.progression.get('bad')
            short_result = 'Bad'
        fig, ax = plt.subplots(figsize=(10, 1))
        ax.barh(0, 100, color=colors[1])
        ax.barh(0, plagiarism_percentage, color=colors[0])
        ax.text(50, 1, f'{plagiarism_percentage}% - {short_result}',
                ha='center', va='center', color='black', fontsize=14)
        ax.set_xlim(0, 100)
        ax.set_ylim(-1, 1)
        ax.axis('off')
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight', pad_inches=0)
        plt.close()
        return buffer

    @staticmethod
    def _validate_percentage(plagiarism_percentage):
        if plagiarism_percentage < 0:
            return 100
        elif plagiarism_percentage > 100:
            return 0
        else:
            originality_percentage = 100 - plagiarism_percentage
            return round(originality_percentage, 1)

    @staticmethod
    def _replace_text_in_paragraphs(paragraphs, old_text, new_text):
        for paragraph in paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

    def _replace_text_in_tables(self, tables, old_text, new_text):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_text_in_paragraphs(cell.paragraphs, old_text, new_text)

    def _replace_text_in_headers(self, doc, old_text, new_text):
        for section in doc.sections:
            header = section.header
            footer = section.footer

            first_page_header = section.first_page_header
            first_page_footer = section.first_page_footer

            for hdr in [header, first_page_header]:
                self._replace_text_in_paragraphs(hdr.paragraphs, old_text, new_text)
                self._replace_text_in_tables(hdr.tables, old_text, new_text)

            for ftr in [footer, first_page_footer]:
                self._replace_text_in_paragraphs(ftr.paragraphs, old_text, new_text)
                self._replace_text_in_tables(ftr.tables, old_text, new_text)

    def _replace_text(self, doc, old_text, new_text, tables=False, headers=False):
        self._replace_text_in_paragraphs(doc.paragraphs, old_text, new_text)
        if tables:
            self._replace_text_in_tables(doc.tables, old_text, new_text)
        if headers:
            self._replace_text_in_headers(doc, old_text, new_text)

    def _add_progress_bar_image(self, doc, plagiarism_percentage):
        buffer = self._generate_progress_bar(self._validate_percentage(plagiarism_percentage))
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(buffer)

    @staticmethod
    def _add_plagiarism_detail(doc, plagiarisms):
        start_index = 1
        for plagiarism in plagiarisms:
            paragraph = doc.add_paragraph(f'{start_index}. {plagiarism.get('sentence')}')
            run = paragraph.runs[0]
            run.font.name = 'Arial'
            paragraph = doc.add_paragraph(f'{plagiarism.get('link')}')
            run = paragraph.runs[0]
            run.font.name = 'Arial'
            start_index += 1

    @staticmethod
    def _save_file(doc, save_path, filename):
        filename = f'{save_path}/report_{filename}.docx'
        if os.path.exists(filename):
            name, ext = os.path.splitext(filename)
            index = 1
            new_filename = "{}_[{}]{}".format(name, index, ext)
            while os.path.exists(new_filename):
                index += 1
                new_filename = "{}_[{}]{}".format(name, index, ext)
            filename = new_filename
        doc.save(filename)

    def generate_document(self, found_plagiarism, plagiarism_percentage, filename, save_path):
        doc = self._load_template()
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self._replace_text(doc, '[time_generated]', current_time, True, True)
        self._replace_text(doc, '[filename]', f'"{filename}"')
        self._add_progress_bar_image(doc, plagiarism_percentage)
        self._add_plagiarism_detail(doc, found_plagiarism)
        self._save_file(doc, save_path, filename)
